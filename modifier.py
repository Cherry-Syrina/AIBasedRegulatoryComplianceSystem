# modifier.py (fixed for Linux/Streamlit Cloud deployment)
import streamlit as st
from docx import Document
from docx.shared import Pt
import tempfile
import os

from llm_helper import call_llm_with_fallback


def _rewrite_clause_with_llm(item):
    """
    Call LLM to rewrite a clause safely.
    """
    clause = item.get("clause", "")
    prompt = (
        "Rewrite the following high-risk contract clause into a safer, "
        "more compliant version while preserving its intent:\n\n"
        f"{clause}"
    )
    return call_llm_with_fallback(prompt)


def modify_contract_docx(analysis_results):
    """
    Create a Word document with High Risk Clauses + AI-rewritten safe versions.
    Returns the path to the generated .docx file.
    """
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx_path = tmp_file.name

    doc = Document()

    doc.add_heading("High Risk Clause Report", level=0)
    doc.add_paragraph(
        "This document summarizes the high-risk clauses identified in the contract "
        "and provides AI-generated safer rewrites for compliance."
    )
    doc.add_page_break()

    for item in analysis_results:
        if item.get("severity", "").lower() == "high":
            original_clause = str(item.get("clause", "")).strip()
            safe_clause = _rewrite_clause_with_llm(item)

            doc.add_heading("High Risk Clause:", level=1)
            para = doc.add_paragraph(original_clause)
            para.style.font.size = Pt(11)

            doc.add_heading("Rewritten Safe Clause:", level=1)
            para = doc.add_paragraph(safe_clause)
            para.style.font.size = Pt(11)

            doc.add_paragraph("-" * 80)

    doc.save(output_docx_path)
    return output_docx_path


def render_download_buttons(analysis_results):
    """
    Generate Word report and render Streamlit download button.
    (PDF conversion removed — not supported on Linux cloud without LibreOffice)
    """
    try:
        docx_path = modify_contract_docx(analysis_results)

        with open(docx_path, "rb") as f:
            st.download_button(
                label="⬇️ Download High Risk Clause Report (Word)",
                data=f,
                file_name="High_Risk_Clause_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        st.error(f"❌ Failed to generate report: {e}")
